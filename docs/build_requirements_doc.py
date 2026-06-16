# -*- coding: utf-8 -*-
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "CarRentalSystem_Functional_Requirements.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B7C3D0"
MUTED = "555555"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        paragraph.style = "Body Table"
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_table_width(table, widths):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def style_table(table, widths=None, header=True):
    if widths:
        set_table_width(table, widths)
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                if p.style.name == "Normal":
                    p.style = "Body Table"
        if header and i == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_BLUE)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table, widths)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_note_table(doc, title, body, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.style = "Body Text"
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(" " + body)
    style_table(table, [6.5], header=False)
    doc.add_paragraph()


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    body_table = styles.add_style("Body Table", 1)
    body_table.base_style = normal
    body_table.font.name = "Calibri"
    body_table._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    body_table.font.size = Pt(9)
    body_table.paragraph_format.space_after = Pt(0)
    body_table.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        s = styles[style_name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title.paragraph_format.space_after = Pt(6)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(12)

    for list_style in ("List Bullet", "List Number"):
        s = styles[list_style]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        s.font.size = Pt(11)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.text = "CarRentalSystem - Functional Requirements"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def add_title_page(doc):
    p = doc.add_paragraph(style="Title")
    p.add_run("CarRentalSystem")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph(style="Subtitle")
    sub.add_run("Tài liệu requirement chức năng, phân quyền và cách hệ thống hoạt động")
    doc.add_paragraph(
        "Phạm vi đọc mã: toàn bộ thư mục src/main/java và src/main/webapp, "
        "kèm đối chiếu cấu trúc SQL hiện có để hiểu nghiệp vụ dữ liệu."
    )

    src_files = sorted(p for p in (ROOT / "src").rglob("*") if p.is_file())
    java_count = len([p for p in src_files if p.suffix == ".java"])
    jsp_count = len([p for p in src_files if p.suffix == ".jsp"])
    meta_rows = [
        ("Dự án", "CarRentalSystem"),
        ("Nền tảng", "Java 17, Jakarta Servlet/JSP, Maven WAR, SQL Server"),
        ("Nguồn đã đọc", f"{len(src_files)} file trong src ({java_count} Java, {jsp_count} JSP, CSS/XML/JSP include)"),
        ("Ngày tạo", datetime.now().strftime("%Y-%m-%d")),
        ("Mục đích", "Làm tài liệu mô tả chức năng chi tiết, role, phân quyền và luồng xử lý"),
    ]
    add_table(doc, ["Thông tin", "Giá trị"], meta_rows, widths=[1.7, 4.8])
    add_note_table(
        doc,
        "Kết luận nhanh.",
        "Dự án đã có đầy đủ khung MVC cơ bản cho thuê xe: tìm xe, đặt xe, xử lý hợp đồng, quản lý xe, gán tài xế, lịch tài xế và quản trị tài khoản. Một số điểm bảo mật và lỗi routing cần xử lý trước khi dùng như sản phẩm thật.",
    )
    doc.add_page_break()


def add_overview(doc):
    doc.add_heading("1. Tổng quan hệ thống", level=1)
    doc.add_paragraph(
        "CarRentalSystem là ứng dụng web cho thuê xe theo mô hình Servlet - JSP - DAO - Model. "
        "Người dùng thao tác qua JSP, controller nhận request theo URL, DAO truy vấn SQL Server, "
        "model mang dữ liệu giữa tầng database và giao diện."
    )
    add_table(
        doc,
        ["Lớp", "Thành phần", "Vai trò"],
        [
            ("Presentation", "JSP trong WEB-INF/views, includes/header.jsp, includes/footer.jsp, css/style.css", "Hiển thị màn hình, form, bảng dữ liệu, nút thao tác và thông báo."),
            ("Controller", "Các Servlet trong com.carrental.controller", "Nhận request, kiểm tra dữ liệu cơ bản, gọi DAO, forward JSP hoặc redirect."),
            ("DAO", "UserDAO, CarDAO, ContractDAO, DriverDAO", "Thao tác SQL Server bằng PreparedStatement và transaction ở nghiệp vụ tạo hợp đồng/user."),
            ("Model", "User, Car, Contract, ContractDetail, Driver, DriverAssignment, Role, Customer, CarType", "POJO ánh xạ bảng và dữ liệu join dùng để hiển thị."),
            ("Filter", "AuthFilter, CharacterEncodingFilter", "Bảo vệ URL cần đăng nhập/phân quyền và ép UTF-8 request/response."),
            ("Database", "CarRentalDB", "Lưu user, role, xe, hợp đồng, chi tiết hợp đồng, tài xế, phân công và lịch sử trạng thái."),
        ],
        widths=[1.2, 2.4, 2.9],
    )

    doc.add_heading("1.1 Phạm vi chức năng chính", level=2)
    add_bullets(
        doc,
        [
            "Khách vãng lai xem danh mục xe, tìm xe còn trống theo thời gian và đăng ký tài khoản.",
            "Customer đặt một hoặc nhiều xe, chọn thuê kèm tài xế, xem/hủy hợp đồng theo điều kiện.",
            "Staff duyệt hoặc xử lý trạng thái hợp đồng trong vòng đời thuê xe.",
            "Manager quản lý đội xe và gán tài xế cho hợp đồng đã duyệt.",
            "Driver xem lịch lái được gán và cập nhật tiến độ chuyến.",
            "Admin quản lý tài khoản, trạng thái tài khoản và role.",
        ],
    )


def add_roles(doc):
    doc.add_heading("2. Role và phân quyền", level=1)
    doc.add_paragraph(
        "Role được lưu trong bảng Roles và User_Roles. Sau khi đăng nhập, LoginServlet lưu User vào session với key loggedInUser và danh sách role vào userRoles. AuthFilter kiểm soát các URL cần đăng nhập."
    )
    add_table(
        doc,
        ["Actor/Role", "URL chính", "Quyền/chức năng", "Ghi chú"],
        [
            ("Guest", "/search, /cars, /login, /register", "Xem trang chủ, tìm xe, xem danh mục xe, đăng ký, đăng nhập.", "Không được vào /book, /my-contracts, /contract-detail và các dashboard."),
            ("CUSTOMER", "/book, /my-contracts, /contract-detail", "Đặt xe, xem danh sách hợp đồng của mình, xem chi tiết hợp đồng, hủy hợp đồng PENDING_REVIEW hoặc ACCEPTED.", "Đăng ký public luôn tạo role CUSTOMER và bản ghi Customers."),
            ("STAFF", "/staff/dashboard, /staff/process", "Xem/lọc hợp đồng, duyệt, từ chối, xác nhận cọc, giao xe, nhận xe trả, hoàn tất thanh toán, hủy hợp đồng.", "Đường link chi tiết trên JSP đang trỏ /staff/contract-detail nhưng servlet hiện không map URL này."),
            ("MANAGER", "/manager/dashboard, /staff/*", "Có quyền staff, thêm/sửa xe, xem hợp đồng đã duyệt, gán tài xế, xem danh sách tài xế đang hoạt động.", "AuthFilter cho MANAGER vào /staff/* và /manager/*."),
            ("DRIVER", "/driver/schedule", "Xem lịch lái, xác nhận nhận xe, bắt đầu chuyến, hoàn tất chuyến.", "Server hiện chưa kiểm tra assignmentId có thuộc chính tài xế đang đăng nhập trong POST cập nhật."),
            ("ADMIN", "/admin/dashboard, /staff/*, /manager/*, /driver/*", "Quản trị tài khoản, role, trạng thái user; cũng được AuthFilter cho vào staff/manager/driver.", "Nếu admin không có hồ sơ Driver thì /driver/schedule hiển thị rỗng."),
        ],
        widths=[1.0, 1.55, 2.65, 1.3],
    )

    doc.add_heading("2.1 Luồng điều hướng sau đăng nhập", level=2)
    add_numbered(
        doc,
        [
            "LoginServlet kiểm tra username/password với UserDAO.login và chỉ nhận user có Status = ACTIVE.",
            "Nếu login thành công, session lưu loggedInUser và userRoles.",
            "Nếu request có redirect bắt đầu bằng '/', hệ thống redirect theo tham số này.",
            "Nếu không có redirect, hệ thống ưu tiên role theo thứ tự ADMIN, MANAGER, STAFF, DRIVER, sau đó mới về /search.",
        ],
    )


def add_functional_requirements(doc):
    doc.add_heading("3. Requirement chức năng chi tiết", level=1)
    requirements = [
        ("FR-01", "Trang chủ và tìm xe", "Người dùng nhập loại xe 4/5/7 chỗ hoặc tất cả, thời gian nhận xe, thời gian trả xe. Hệ thống kiểm tra ngày trả phải sau ngày nhận và trả danh sách xe còn trống.", "CarSearchServlet, CarDAO.findAvailableCars, search.jsp"),
        ("FR-02", "Danh mục xe", "Người dùng xem toàn bộ đội xe, lọc theo số chỗ, hãng xe, trạng thái. Xe AVAILABLE có nút đặt xe mở modal chọn ngày.", "CarCatalogServlet, car-catalog.jsp"),
        ("FR-03", "Đăng ký", "Người dùng tạo tài khoản bằng username, email, password, họ tên, phone, địa chỉ. Hệ thống kiểm tra trùng username/email, tạo Users, gán role CUSTOMER và tạo Customers.", "RegisterServlet, UserDAO.registerCustomer"),
        ("FR-04", "Đăng nhập/đăng xuất", "Người dùng đăng nhập bằng username/password. Hệ thống tạo session, lưu role, redirect theo role hoặc redirect cũ. Logout hủy session.", "LoginServlet, LogoutServlet, AuthFilter"),
        ("FR-05", "Đặt xe", "Customer chọn một hoặc nhiều xe, xác nhận thời gian, chọn thuê kèm tài xế từng xe, nhập địa điểm nhận/trả. Hệ thống tính ngày thuê tối thiểu 1 ngày, tiền cọc, tiền thuê xe và phí tài xế cố định 300,000 VND/ngày.", "BookingServlet, ContractDAO.createContractWithDetails, booking.jsp"),
        ("FR-06", "Hợp đồng của tôi", "Customer xem hợp đồng của chính mình, trạng thái, tiền cọc, tổng tiền, ngày tạo. Có nút xem chi tiết và hủy nếu trạng thái còn cho phép.", "CustomerContractsServlet, my-contracts.jsp"),
        ("FR-07", "Chi tiết hợp đồng khách hàng", "Customer xem thông tin thời gian, địa điểm, chi phí, từng xe thuê, phí tài xế, tài xế được gán và timeline trạng thái.", "ContractDetailServlet, contract-detail.jsp, DriverDAO.getAssignmentByDetailId"),
        ("FR-08", "Xử lý hợp đồng staff", "Staff/Manager/Admin xem toàn bộ hợp đồng, lọc theo trạng thái và cập nhật trạng thái: accept, reject, deposit_paid, car_picked_up, car_returned, final_payment, cancel.", "StaffDashboardServlet, ContractProcessServlet, ContractDAO.updateContractStatus"),
        ("FR-09", "Quản lý xe", "Manager/Admin xem danh sách xe, thêm xe, sửa xe với biển số, hãng, model, loại xe, năm, màu, hộp số, nhiên liệu, ODO, giá/ngày, cọc, trạng thái, image URL, mô tả.", "ManagerDashboardServlet, CarDAO.insertCar/updateCar, manager-dashboard.jsp"),
        ("FR-10", "Gán tài xế", "Manager/Admin chọn hợp đồng ACCEPTED, xem chi tiết xe cần tài xế, kiểm tra xe đã có tài xế hay chưa, xem tài xế rảnh/bận theo khoảng thời gian và gán tài xế.", "ManagerDashboardServlet, DriverDAO.getDriversWithAvailability, DriverDAO.assignDriver"),
        ("FR-11", "Lịch tài xế", "Driver xem lịch được gán, thông tin xe, khách, địa điểm, thời gian. Driver cập nhật trạng thái ASSIGNED, HANDOVER_RECEIVED, TRIP_IN_PROGRESS, TRIP_COMPLETED.", "DriverScheduleServlet, DriverDAO.getScheduleByDriverId/updateAssignmentStatus"),
        ("FR-12", "Quản trị tài khoản", "Admin xem danh sách user, tạo user mới với role, đổi trạng thái ACTIVE/LOCKED/DISABLED và cập nhật role.", "AdminDashboardServlet, UserDAO.createUser/updateUserStatus/updateUserRoles"),
        ("FR-13", "Lỗi và session", "web.xml cấu hình session timeout 30 phút, error page chung cho 403/404/500, welcome file index.jsp redirect về /search.", "web.xml, index.jsp, error.jsp"),
    ]
    add_table(doc, ["ID", "Module", "Requirement", "Nguồn code"], requirements, widths=[0.65, 1.25, 3.6, 1.0])


def add_workflows(doc):
    doc.add_heading("4. Cách hệ thống hoạt động theo luồng", level=1)
    workflows = [
        (
            "4.1 Luồng tìm xe và đặt xe",
            [
                "Guest hoặc Customer vào /search hoặc /cars.",
                "Ở /search, hệ thống gọi CarDAO.findAvailableCars để loại xe không AVAILABLE, xe trùng hợp đồng đang hoạt động và xe đang bảo trì.",
                "Người dùng chọn xe và chuyển sang /book. AuthFilter yêu cầu đăng nhập trước khi tiếp tục.",
                "BookingServlet.doGet lấy danh sách carId, tính số ngày thuê theo giờ thuê / 24 và làm tròn lên, tối thiểu 1 ngày.",
                "BookingServlet.doPost kiểm tra user có CustomerID, tạo Contract và các Contract_Details trong một transaction.",
                "Contract ban đầu lấy status mặc định PENDING_REVIEW từ database; Contract_Details mặc định BOOKED.",
            ],
        ),
        (
            "4.2 Luồng xử lý hợp đồng",
            [
                "Staff xem /staff/dashboard và lọc hợp đồng theo status nếu cần.",
                "Với PENDING_REVIEW: staff có thể accept hoặc reject.",
                "Với ACCEPTED: staff xác nhận deposit_paid.",
                "Với DEPOSIT_PAID: staff xác nhận car_picked_up.",
                "Với CAR_PICKED_UP: staff xác nhận car_returned.",
                "Với CAR_RETURNED: staff xác nhận final_payment để chuyển FINAL_PAYMENT_COMPLETED.",
                "Hệ thống ghi Contract_Status_History khi đổi trạng thái và đồng bộ DetailStatus theo một số trạng thái chính.",
            ],
        ),
        (
            "4.3 Luồng gán tài xế",
            [
                "Manager mở tab Gán tài xế trong /manager/dashboard?tab=assign.",
                "Hệ thống chỉ liệt kê hợp đồng trạng thái ACCEPTED để chọn.",
                "Với từng ContractDetail requiresDriver = true, hệ thống kiểm tra đã có Driver_Assignments active chưa.",
                "Nếu chưa có, danh sách tài xế được hiển thị với trạng thái rảnh/bận dựa trên lịch trùng khoảng pickup-return.",
                "Khi gán, DriverDAO.insert Driver_Assignments. Trigger SQL sẽ từ chối nếu trùng lịch hoặc detail không yêu cầu tài xế.",
            ],
        ),
        (
            "4.4 Luồng tài xế",
            [
                "Driver đăng nhập sẽ được điều hướng tới /driver/schedule nếu role DRIVER là role ưu tiên cao nhất của user.",
                "DriverScheduleServlet lấy DriverID từ UserID, sau đó lấy lịch từ DriverDAO.getScheduleByDriverId.",
                "Driver bấm nút cập nhật trạng thái theo tiến trình: ASSIGNED -> HANDOVER_RECEIVED -> TRIP_IN_PROGRESS -> TRIP_COMPLETED.",
                "Khi chuyển HANDOVER_RECEIVED hoặc TRIP_COMPLETED, DAO lưu thêm thời điểm tương ứng.",
            ],
        ),
        (
            "4.5 Luồng admin",
            [
                "Admin vào /admin/dashboard để xem toàn bộ users và role của từng user.",
                "Admin tạo user mới. UserDAO.createUser vừa thêm Users/User_Roles vừa tạo Customers/Drivers/Employees tùy role ban đầu.",
                "Admin cập nhật trạng thái user để khóa hoặc vô hiệu hóa tài khoản.",
                "Admin cập nhật role bằng cách xóa toàn bộ role cũ rồi thêm role mới.",
            ],
        ),
    ]
    for heading, steps in workflows:
        doc.add_heading(heading, level=2)
        add_numbered(doc, steps)

    doc.add_heading("4.6 Vòng đời trạng thái", level=2)
    add_table(
        doc,
        ["Đối tượng", "Trạng thái", "Ý nghĩa trong hệ thống"],
        [
            ("Contract", "PENDING_REVIEW", "Hợp đồng mới được customer gửi, chờ staff duyệt."),
            ("Contract", "ACCEPTED", "Staff đã duyệt. Manager có thể gán tài xế nếu có xe cần tài xế."),
            ("Contract", "REJECTED", "Staff từ chối hợp đồng."),
            ("Contract", "DEPOSIT_PAID", "Đã ghi nhận đặt cọc."),
            ("Contract", "CAR_PICKED_UP", "Khách đã nhận xe."),
            ("Contract", "CAR_RETURNED", "Khách đã trả xe."),
            ("Contract", "FINAL_PAYMENT_COMPLETED", "Hoàn tất thanh toán cuối cùng."),
            ("Contract", "CANCELLED", "Hợp đồng bị hủy."),
            ("ContractDetail", "BOOKED, PICKED_UP, RETURNED, CANCELLED", "Trạng thái từng dòng xe thuê, được map một phần từ trạng thái hợp đồng."),
            ("DriverAssignment", "ASSIGNED, HANDOVER_RECEIVED, TRIP_IN_PROGRESS, TRIP_COMPLETED, CANCELLED", "Trạng thái công việc của tài xế."),
            ("Car", "AVAILABLE, MAINTENANCE, INACTIVE, RETIRED", "Trạng thái xe dùng cho tìm kiếm, quản lý và catalog."),
            ("User", "ACTIVE, LOCKED, DISABLED", "Trạng thái tài khoản. Login chỉ cho ACTIVE."),
        ],
        widths=[1.25, 2.15, 3.1],
    )


def add_data_model(doc):
    doc.add_heading("5. Dữ liệu và DAO", level=1)
    doc.add_paragraph(
        "Các DAO đều dùng PreparedStatement, giúp giảm nguy cơ SQL injection ở phần truy vấn chính. Một số nghiệp vụ dùng transaction để đảm bảo tạo dữ liệu đồng bộ."
    )
    add_table(
        doc,
        ["Bảng/Model", "Vai trò nghiệp vụ", "DAO chính"],
        [
            ("Users, Roles, User_Roles", "Tài khoản, phân quyền, trạng thái đăng nhập.", "UserDAO"),
            ("Customers", "Hồ sơ customer dùng để tạo và lọc hợp đồng của khách.", "UserDAO, ContractDAO"),
            ("Employees", "Hồ sơ nhân sự cho STAFF/MANAGER/ADMIN khi admin tạo user.", "UserDAO"),
            ("Drivers", "Hồ sơ tài xế, bằng lái, phí cơ bản, trạng thái hoạt động.", "UserDAO, DriverDAO"),
            ("Car_Types, Cars", "Loại xe, thông tin xe, giá thuê, cọc, trạng thái.", "CarDAO"),
            ("Contracts", "Hợp đồng thuê xe tổng, thời gian, địa điểm, tiền cọc, tổng tiền, trạng thái.", "ContractDAO"),
            ("Contract_Details", "Từng xe trong hợp đồng, phí xe, phí tài xế, số ngày, line total.", "ContractDAO"),
            ("Driver_Assignments", "Tài xế được gán cho từng chi tiết hợp đồng.", "DriverDAO"),
            ("Car_Maintenance", "Lịch bảo trì dùng để loại xe khỏi tìm kiếm.", "CarDAO.findAvailableCars"),
            ("Contract_Status_History", "Lịch sử đổi trạng thái hợp đồng.", "ContractDAO.updateContractStatus"),
        ],
        widths=[1.8, 3.2, 1.5],
    )

    doc.add_heading("5.1 Mapping Servlet - JSP - DAO", level=2)
    add_table(
        doc,
        ["URL", "Servlet", "JSP", "DAO/Method chính"],
        [
            ("/search", "CarSearchServlet", "search.jsp", "CarDAO.findAvailableCars"),
            ("/cars", "CarCatalogServlet", "car-catalog.jsp", "CarDAO.getAllCars"),
            ("/login", "LoginServlet", "login.jsp", "UserDAO.login, getUserRoles"),
            ("/register", "RegisterServlet", "register.jsp", "UserDAO.registerCustomer"),
            ("/book", "BookingServlet", "booking.jsp", "CarDAO.getCarById, ContractDAO.createContractWithDetails"),
            ("/my-contracts", "CustomerContractsServlet", "my-contracts.jsp", "ContractDAO.getContractsByCustomerId, updateContractStatus"),
            ("/contract-detail", "ContractDetailServlet", "contract-detail.jsp", "ContractDAO.getContractById, DriverDAO.getAssignmentByDetailId"),
            ("/staff/dashboard", "StaffDashboardServlet", "staff-dashboard.jsp", "ContractDAO.getAllContracts/getContractsByStatus"),
            ("/staff/process", "ContractProcessServlet", "redirect", "ContractDAO.updateContractStatus"),
            ("/manager/dashboard", "ManagerDashboardServlet", "manager-dashboard.jsp", "CarDAO, ContractDAO, DriverDAO"),
            ("/driver/schedule", "DriverScheduleServlet", "driver-schedule.jsp", "DriverDAO.getScheduleByDriverId/updateAssignmentStatus"),
            ("/admin/dashboard", "AdminDashboardServlet", "admin-dashboard.jsp", "UserDAO.getAllUsers/createUser/updateUserStatus/updateUserRoles"),
            ("/logout", "LogoutServlet", "redirect", "Session invalidate"),
        ],
        widths=[1.2, 1.7, 1.55, 2.05],
    )


def add_nonfunctional_and_gaps(doc):
    doc.add_heading("6. Kiểm tra hiện trạng và vấn đề cần xử lý", level=1)
    add_note_table(
        doc,
        "Ghi chú kiểm tra build.",
        "Trong môi trường hiện tại, lệnh mvn không có trong PATH nên không chạy được Maven test/package. Project có target/CarRentalSystem.war đã được tạo trước đó. Khi thử javac thủ công, mã biên dịch được đến lúc compiler bị chặn quyền khi đóng JAR servlet API ngoài workspace, nên tài liệu này dựa trên đọc source và đối chiếu artifact sẵn có.",
    )
    gaps = [
        ("P1", "Password người dùng lưu và so sánh dạng plain text dù field tên PasswordHash.", "Rủi ro bảo mật rất cao nếu DB bị lộ.", "Dùng BCrypt/Argon2, migrate dữ liệu test, đổi UserDAO.login/register/createUser."),
        ("P1", "Staff dashboard link tới /staff/contract-detail nhưng servlet hiện map /contract-detail và chỉ cho customer owner xem.", "Staff bấm mã hợp đồng có thể bị 404 hoặc không xem được chi tiết.", "Tạo StaffContractDetailServlet riêng hoặc sửa link/logic phân quyền chi tiết hợp đồng cho staff."),
        ("P2", "Không có CSRF token cho các form POST.", "User đã đăng nhập có thể bị ép gửi action ngoài ý muốn.", "Thêm CSRF token trong session và validate ở các servlet POST."),
        ("P2", "DriverScheduleServlet.doPost chỉ nhận assignmentId và status, không kiểm tra assignment thuộc driver đăng nhập.", "Driver có thể cập nhật assignment khác nếu biết ID.", "Update query nên JOIN theo DriverID của user hiện tại và validate trạng thái hợp lệ."),
        ("P2", "ContractProcessServlet cho phép đổi trạng thái theo action nhưng server không kiểm tra trạng thái hiện tại.", "Request thủ công có thể nhảy bước workflow.", "Áp dụng state machine ở server: chỉ cho transition hợp lệ từ oldStatus."),
        ("P2", "Ghi chú my-contracts.jsp nói có thể hủy DEPOSIT_PAID, nhưng servlet chỉ cho PENDING_REVIEW và ACCEPTED.", "Người dùng bị hiểu sai điều kiện hủy.", "Chọn rule chính thức rồi sửa cả JSP và CANCELLABLE trong servlet."),
        ("P2", "Admin updateRoles chỉ đổi User_Roles, không tạo/xóa Customers/Drivers/Employees tương ứng.", "User được thêm role sau này có thể thiếu hồ sơ phụ, gây lỗi đặt xe/lịch tài xế.", "Đồng bộ profile tables khi cập nhật role hoặc khóa không cho đổi role yêu cầu profile."),
        ("P2", "Các parse số/ngày dùng Integer.parseInt, Long.parseLong, LocalDateTime.parse, BigDecimal trực tiếp.", "Input sai hoặc request thủ công có thể gây 500.", "Thêm try/catch, validate form server-side, trả error thân thiện."),
        ("P2", "BookingServlet không kiểm tra returnAt > pickupAt trong /book GET/POST.", "Request trực tiếp có thể đến DB constraint rồi fail khó hiểu.", "Validate lại ngày ở BookingServlet, không chỉ ở SearchServlet/JS."),
        ("P3", "Login redirect nhận mọi chuỗi bắt đầu bằng '/'.", "Có thể gây redirect không mong muốn nếu bị truyền tham số thủ công.", "Chỉ cho redirect nội bộ cùng contextPath và chặn chuỗi bắt đầu bằng '//'."),
        ("P3", "Chưa có test tự động trong src/test.", "Khó đảm bảo các luồng chính không regress.", "Thêm unit test DAO/service nếu tách service, và integration test servlet cơ bản."),
    ]
    add_table(doc, ["Ưu tiên", "Vấn đề", "Ảnh hưởng", "Đề xuất"], gaps, widths=[0.65, 2.05, 1.85, 1.95])

    doc.add_heading("6.1 Checklist trước khi demo/nộp", level=2)
    add_bullets(
        doc,
        [
            "Sửa URL chi tiết hợp đồng cho staff hoặc tạo servlet chi tiết hợp đồng dành cho staff.",
            "Chốt rule hủy hợp đồng và đồng bộ text giao diện với server.",
            "Ẩn thông tin mật khẩu database khỏi source trước khi chia sẻ repository.",
            "Đổi plain text password sang hash nếu yêu cầu bảo mật được chấm.",
            "Chạy Maven package trên máy có Maven và Tomcat 10.1/Jakarta Servlet 6.",
            "Kiểm thử thủ công các tài khoản CUSTOMER, STAFF, MANAGER, DRIVER, ADMIN theo sql/setup-database.sql.",
        ],
    )


def add_appendix(doc):
    doc.add_heading("7. Phụ lục: file source đã đọc", level=1)
    src_files = sorted(p.relative_to(ROOT).as_posix() for p in (ROOT / "src").rglob("*") if p.is_file())
    rows = []
    for idx, path in enumerate(src_files, 1):
        group = "Java" if path.endswith(".java") else "JSP" if path.endswith(".jsp") else "Web/CSS/XML"
        rows.append((str(idx), group, path))
    add_table(doc, ["#", "Nhóm", "File"], rows, widths=[0.5, 1.1, 4.9])


def build():
    doc = Document()
    setup_styles(doc)
    add_title_page(doc)
    add_overview(doc)
    add_roles(doc)
    add_functional_requirements(doc)
    add_workflows(doc)
    add_data_model(doc)
    add_nonfunctional_and_gaps(doc)
    add_appendix(doc)
    doc.core_properties.title = "CarRentalSystem Functional Requirements"
    doc.core_properties.subject = "Requirement, phân quyền, chức năng và luồng hoạt động"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "CarRentalSystem, PRJ301, requirements, role, servlet, JSP"
    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
